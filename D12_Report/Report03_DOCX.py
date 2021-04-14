# pip install python-docx
from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Bảng điểm', 0) # level=0: Title
document.add_paragraph('Họ tên: Trần Anh', style='Heading 3')
document.add_paragraph('Python', style='List Bullet')
document.add_paragraph('SQL', style='List Bullet')
document.add_paragraph('Visualization', style='List Bullet')
document.add_paragraph('HTML5', style='List Number')
document.add_paragraph('CSS3', style='List Number')
document.add_heading('Chi tiết', level=1)

records = [
    { "tenmon" : "Toán", "diem": 9, "sotc": 3 },
    { "tenmon" : "Anh văn", "diem": 7, "sotc": 3 },
]

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Tên môn'
hdr_cells[1].text = 'Số tín chỉ'
hdr_cells[2].text = 'Điểm'
for mon in records:
    row_cells = table.add_row().cells
    row_cells[0].text = mon["tenmon"]
    row_cells[1].text = str(mon["sotc"])
    row_cells[2].text = str(mon["diem"])

document.save('BangDiem.docx')

# pip install docx2pdf
from docx2pdf import convert
convert("BangDiem.docx", "BangDiem.pdf")